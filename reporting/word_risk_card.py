import os
import tempfile
from datetime import datetime

import numpy as np
import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

from analytics.risk_metrics import compute_empirical_risk
from analytics.risk_extra import compute_basic_risk_snapshot_main
from analytics.risk_utils import returns, to_simple


# ---------------------------------------------------------------------
# Styl wykresów – kolorowy, ale minimalistyczny
# ---------------------------------------------------------------------
sns.set_style("whitegrid")
sns.set_context("notebook", font_scale=0.9)  # mniejszy ogólny rozmiar niż "talk"

plt.rcParams["font.family"] = "sans-serif"
plt.rcParams["font.sans-serif"] = ["Lato", "DejaVu Sans", "Arial"]

# Kolory osi
plt.rcParams["axes.edgecolor"] = "#333333"
plt.rcParams["axes.labelcolor"] = "#333333"
plt.rcParams["xtick.color"] = "#333333"
plt.rcParams["ytick.color"] = "#333333"

# Rozmiary czcionek na wykresach
plt.rcParams["axes.titlesize"] = 12     # tytuł wykresu
plt.rcParams["axes.labelsize"] = 10     # opisy osi
plt.rcParams["xtick.labelsize"] = 8     # wartości na osi X
plt.rcParams["ytick.labelsize"] = 8     # wartości na osi Y


# ---------------------------------------------------------------------
# Formatowanie liczb i pomocnicze narzędzia
# ---------------------------------------------------------------------
def fmt_pln(value, decimals: int = 2) -> str:
    """Format liczby w stylu polskim z walutą PLN."""
    if value is None:
        return "-"
    try:
        v = float(value)
    except (TypeError, ValueError):
        return "-"
    txt = f"{v:,.{decimals}f}"
    txt = txt.replace(",", " ").replace(".", ",")
    return f"{txt} PLN"


def fmt_dec(value, decimals: int = 4) -> str:
    """Format zwykłej liczby (np. Sharpe) z przecinkiem."""
    if value is None:
        return "-"
    try:
        v = float(value)
    except (TypeError, ValueError):
        return "-"
    txt = f"{v:.{decimals}f}".replace(".", ",")
    return txt


def fmt_pct(value, decimals: int = 2) -> str:
    """Format procentu (np. MDD)."""
    if value is None:
        return "-"
    try:
        v = float(value)
    except (TypeError, ValueError):
        return "-"
    txt = f"{100 * v:.{decimals}f}".replace(".", ",")
    return f"{txt} %"


def _set_document_style_lato(doc: Document):
    """Ustaw font Lato dla podstawowych styli Worda."""
    base_styles = ["Normal", "Heading 1", "Heading 2", "Heading 3"]
    for name in base_styles:
        if name in doc.styles:
            style = doc.styles[name]
            font = style.font
            font.name = "Lato"
            font.size = Pt(11) if name == "Normal" else Pt(16)
            # żeby Word w pełni łapał font
            rfonts = font.element.rPr.rFonts
            rfonts.set(qn("w:eastAsia"), "Lato")
            rfonts.set(qn("w:ascii"), "Lato")
            rfonts.set(qn("w:hAnsi"), "Lato")


# ---------------------------------------------------------------------
# Pomocnicze funkcje do wykresów
# ---------------------------------------------------------------------
def _save_plot(fig):
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    fig.savefig(tmp.name, bbox_inches="tight", dpi=200)
    plt.close(fig)
    return tmp.name


def _format_dates(ax):
    """Ładne daty na osi X, bez rozjeżdżania się etykiet."""
    ax.xaxis.set_major_locator(mdates.AutoDateLocator())
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m"))
    fig = ax.get_figure()
    fig.autofmt_xdate()


def _get_palette():
    return sns.color_palette("tab10")


def _plot_synthetic_nav(port_rets: pd.Series):
    cum_nav = (1 + port_rets).cumprod()
    fig, ax = plt.subplots(figsize=(6, 3))
    color = _get_palette()[0]
    ax.plot(cum_nav.index, cum_nav.values, color=color, linewidth=2)
    ax.set_title("Synthetic NAV")
    ax.set_xlabel("")
    ax.set_ylabel("NAV")
    _format_dates(ax)
    sns.despine()
    return _save_plot(fig)


def _plot_rolling_vol(port_rets: pd.Series, window: int = 60, trading_days: int = 252):
    rolling_vol = port_rets.rolling(window).std() * np.sqrt(trading_days)
    fig, ax = plt.subplots(figsize=(6, 3))
    color = _get_palette()[1]
    ax.plot(rolling_vol.index, rolling_vol.values, color=color, linewidth=2)
    ax.set_title(f"Rolling {window}-Day Volatility (annualized)")
    ax.set_ylabel("Volatility")
    _format_dates(ax)
    sns.despine()
    return _save_plot(fig)


def _plot_rolling_var(port_rets: pd.Series, window: int = 60, alpha: float = 0.01):
    def _var(x):
        return np.quantile(x, alpha)

    rolling_var = port_rets.rolling(window).apply(_var, raw=True)
    fig, ax = plt.subplots(figsize=(6, 3))
    color = _get_palette()[2]
    ax.plot(rolling_var.index, rolling_var.values, color=color, linewidth=2)
    ax.set_title(f"Rolling {window}-Day VaR (1-alpha={1 - alpha:.0%})")
    ax.set_ylabel("VaR (ret.)")
    _format_dates(ax)
    sns.despine()
    return _save_plot(fig)


def _plot_bar_series(series: pd.Series, title: str, xlabel: str = ""):
    series = series.dropna().sort_values(ascending=False).head(10)
    fig, ax = plt.subplots(figsize=(6, 3))
    color = _get_palette()[3]
    sns.barplot(x=series.values, y=series.index, color=color, ax=ax)
    ax.set_title(title)
    ax.set_xlabel(xlabel)
    sns.despine(left=True, bottom=True)
    return _save_plot(fig)


def _plot_histogram(rets: pd.Series, title: str = "Histogram of Daily Returns"):
    fig, ax = plt.subplots(figsize=(6, 3))
    color = _get_palette()[4]
    sns.histplot(rets.dropna(), color=color, kde=False, bins=30, ax=ax)
    ax.set_title(title)
    ax.set_xlabel("Daily return")
    sns.despine()
    return _save_plot(fig)


# ---------------------------------------------------------------------
# KARTA RYZYKA PORTFELA
# ---------------------------------------------------------------------
class PortfolioRiskCardGenerator:
    def __init__(self, prices: pd.DataFrame, holdings: pd.Series, cash: float, cfg: dict,
                 start_date=None, end_date=None):
        self.prices = prices
        self.holdings = holdings
        self.cash = float(cash)
        self.cfg = cfg

        # zakres dat do podpisu na karcie
        if start_date is not None:
            self.start_date = pd.to_datetime(start_date).date()
        else:
            self.start_date = pd.to_datetime(prices.index.min()).date()

        if end_date is not None:
            self.end_date = pd.to_datetime(end_date).date()
        else:
            self.end_date = pd.to_datetime(prices.index.max()).date()

    def _compute_risk(self):
        """Spójne metryki jak w Summary: compute_empirical_risk + compute_basic_risk_snapshot_main."""
        var_h = int(self.cfg.get("var_horizon_days", 20))
        trading_days = int(self.cfg.get("trading_days", 252))
        risk_window_days = int(self.cfg.get("risk_window_days", 252))
        use_log = bool(self.cfg.get("use_log_returns", True))
        var_conf = float(self.cfg.get("var_confidence", 0.99))
        r_f = float(self.cfg.get("risk_free_rate", 0.0))

        risk_emp = compute_empirical_risk(
            prices=self.prices,
            holdings=self.holdings,
            horizon_days=var_h,
            trading_days=trading_days,
            risk_window_days=risk_window_days,
            use_log_returns=use_log,
            confidence=var_conf,
        )

        risk_extra = compute_basic_risk_snapshot_main(
            prices=self.prices,
            holdings=self.holdings,
            cash=self.cash,
            risk_window_days=risk_window_days,
            var_conf=var_conf,
            var_h=var_h,
            trading_days=trading_days,
            r_f=r_f,
            benchmark=None,  # na razie bez benchmarku
        )

        risk_emp.update(risk_extra)
        return risk_emp

    def generate(self, output_path: str | None = None):
        risk = self._compute_risk()

        use_log = bool(self.cfg.get("use_log_returns", True))
        trading_days = int(self.cfg.get("trading_days", 252))
        var_conf = float(self.cfg.get("var_confidence", 0.99))
        var_h = int(self.cfg.get("var_horizon_days", 20))

        # --- portfelowe zwroty do wykresów ---
        rets = returns(self.prices, log=use_log)
        # wartości spółek
        last = self.prices.ffill().iloc[-1]
        equity_value = float((self.holdings * last).sum())
        if equity_value > 0:
            weights = (self.holdings * last) / equity_value
        else:
            weights = pd.Series(0.0, index=self.prices.columns)
        port_rets = rets.mul(weights, axis=1).sum(axis=1)
        port_rets = to_simple(port_rets, use_log)

        # synthetic NAV / rolling metrics
        nav_plot = _plot_synthetic_nav(port_rets)
        vol_plot = _plot_rolling_vol(port_rets, window=60, trading_days=trading_days)
        var_plot = _plot_rolling_var(port_rets, window=60, alpha=1 - var_conf)

        # wagi do Top10
        weights_map = weights.sort_values(ascending=False)
        weights_plot = _plot_bar_series(weights_map, "Top 10 Holdings", xlabel="Portfolio weight")

        # -----------------------------------------------------------------
        # Dokument Word
        # -----------------------------------------------------------------
        doc = Document()
        _set_document_style_lato(doc)

        title = f"Portfolio Risk Card – {self.end_date.isoformat()}"
        doc.add_heading(title, level=0)

        period_par = doc.add_paragraph(
            f"Okres analizy: {self.start_date.strftime('%d.%m.%Y')} – {self.end_date.strftime('%d.%m.%Y')}"
        )
        period_par.runs[0].font.bold = True

        doc.add_paragraph("")

        # Tabela głównych wskaźników + miejsce na komentarz
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        def add_row(label, value):
            row = table.add_row().cells
            row[0].text = label
            row[1].text = value

        add_row("NAV", fmt_pln(risk.get("nav")))
        add_row("Gotówka", fmt_pln(self.cash))
        add_row("Zmienność dzienna (σ)", fmt_dec(risk.get("daily_vol")))
        add_row("Zmienność roczna (σ)", fmt_dec(risk.get("annual_vol")))
        add_row(f"VaR 1D ({var_conf:.0%})", fmt_pln(risk.get("var_1d")))
        add_row("ES 1D", fmt_pln(risk.get("es_1d")))
        add_row(f"VaR √h, h={var_h}", fmt_pln(risk.get("var_h")))
        add_row(f"ES √h, h={var_h}", fmt_pln(risk.get("es_h")))
        add_row("Max Drawdown", fmt_pct(risk.get("max_drawdown")))
        add_row("Sharpe", fmt_dec(risk.get("sharpe")))
        add_row("Tracking Error", fmt_dec(risk.get("tracking_error")))
        add_row("Information Ratio", fmt_dec(risk.get("information_ratio")))
        add_row("Beta", fmt_dec(risk.get("beta")))


        doc.add_paragraph("")

        # Strona 1 – NAV + vol
        doc.add_picture(nav_plot, width=Inches(6))
        doc.add_picture(vol_plot, width=Inches(6))

        # Strona 2 – VaR + Top10
        doc.add_page_break()
        doc.add_picture(var_plot, width=Inches(6))
        doc.add_picture(weights_plot, width=Inches(6))

        if output_path is None:
            today_str = self.end_date.isoformat()
            output_path = f"output/risk_cards/portfolio_risk_card_{today_str}.docx"

        doc.save(output_path)
        print(f"[OK] Portfolio Risk Card saved to {output_path}")


# ---------------------------------------------------------------------
# KARTA RYZYKA SPÓŁKI
# ---------------------------------------------------------------------
class StockRiskCardGenerator:
    def __init__(self, prices: pd.DataFrame, cfg: dict):
        self.prices = prices
        self.cfg = cfg

    def generate_for_ticker(self, ticker: str, output_path: str | None = None):
        if ticker not in self.prices.columns:
            print(f"[WARN] {ticker} not found in prices")
            return

        series = self.prices[ticker].dropna()
        if series.empty:
            print(f"[WARN] Brak danych cenowych dla {ticker}")
            return

        use_log = bool(self.cfg.get("use_log_returns", True))
        trading_days = int(self.cfg.get("trading_days", 252))
        var_conf = float(self.cfg.get("var_confidence", 0.99))

        # zakres dat z danych spółki
        start_date = pd.to_datetime(series.index.min()).date()
        end_date = pd.to_datetime(series.index.max()).date()

        # zwroty spółki
        rets = returns(self.prices[[ticker]], log=use_log)[ticker]
        rets_simple = to_simple(rets, use_log)

        # podstawowe miary
        daily_vol = float(rets_simple.std())
        annual_vol = daily_vol * np.sqrt(trading_days)

        alpha = 1 - var_conf
        var_1d_ret = float(rets_simple.quantile(alpha))
        tail = rets_simple[rets_simple <= var_1d_ret]
        es_1d_ret = float(tail.mean()) if len(tail) else var_1d_ret

        # wykres ceny
        fig, ax = plt.subplots(figsize=(6, 3))
        color_price = _get_palette()[0]
        ax.plot(series.index, series.values, color=color_price, linewidth=2)
        ax.set_title(f"{ticker} – Price History")
        ax.set_ylabel("Price")
        _format_dates(ax)
        sns.despine()
        price_plot = _save_plot(fig)

        # rolling vol (na zwrotach spółki)
        vol_plot = _plot_rolling_vol(rets_simple, window=60, trading_days=trading_days)
        # histogram zwrotów
        hist_plot = _plot_histogram(rets_simple, title=f"{ticker} – Histogram of Daily Returns")

        # dokument
        doc = Document()
        _set_document_style_lato(doc)

        doc.add_heading(f"Stock Risk Card – {ticker}", level=0)
        period_par = doc.add_paragraph(
            f"Okres analizy: {start_date.strftime('%d.%m.%Y')} – {end_date.strftime('%d.%m.%Y')}"
        )
        period_par.runs[0].font.bold = True
        doc.add_paragraph("")

        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        def add_row(label, value):
            row = table.add_row().cells
            row[0].text = label
            row[1].text = value

        add_row("Ostatnia cena", fmt_pln(series.iloc[-1]))
        add_row("Zmienność dzienna (σ)", fmt_dec(daily_vol))
        add_row("Zmienność roczna (σ)", fmt_dec(annual_vol))
        add_row(f"VaR 1D ({var_conf:.0%})", fmt_pct(-var_1d_ret))  # ujemny VaR jako dodatni %
        add_row("ES 1D", fmt_pct(-es_1d_ret))

        doc.add_paragraph("")
        doc.add_picture(price_plot, width=Inches(6))
        doc.add_picture(vol_plot, width=Inches(6))

        doc.add_page_break()
        doc.add_picture(hist_plot, width=Inches(6))

        if output_path is None:
            today_str = end_date.isoformat()
            output_path = f"output/risk_cards/stock_{ticker}_risk_card_{today_str}.docx"

        doc.save(output_path)
        print(f"[OK] Stock Risk Card for {ticker} saved to {output_path}")
